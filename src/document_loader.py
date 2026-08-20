"""
Load documents and extract text into a common page-based format.

All loaders return the same structure:
    [{"page_number": 1, "text": "..."}, ...]

PDF  → one record per PDF page
DOCX → paragraphs grouped into pseudo-pages (~1500 chars each)

Why keep page numbers?
  When retrieval returns a chunk, knowing its page/section helps you verify
  the result and display source citations.
"""

import io
from pathlib import Path
from typing import Union

import fitz  # PyMuPDF

# Target size for grouping DOCX paragraphs into pseudo-pages.
_DOCX_PAGE_TARGET_CHARS = 1500

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def load_pdf(pdf_path: Union[str, Path]) -> list[dict]:    #Create a function called load_pdf that accepts a file path and returns a list of dictionaries.
    """Extract text from every page of a PDF file on disk."""
    pdf_path = Path(pdf_path)   ## convert the file path to a Path object.- path ma convert karse
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    with open(pdf_path, "rb") as f:    #open means open the file in binary mode.,  rb--read in binary mode.
        return load_pdf_from_bytes(f.read())
 

def load_pdf_from_bytes(pdf_bytes: bytes) -> list[dict]:   #this creates func that accepts raw Pdf bytes  than file name  (why not filename- as u dont have to save it fitst)
    """Extract text from a PDF provided as raw bytes."""
    if not pdf_bytes:    #agar bytes empty hai to error throw karega.
        raise ValueError("PDF bytes are empty.")

    pages: list[dict] = []

    with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
        if doc.page_count == 0:
            raise ValueError("PDF has no pages.")

        for page_index in range(doc.page_count):
            page = doc[page_index]
            text = page.get_text("text").strip()   ##for eg. it removes whitespaces and newlines.
            pages.append(
                {
                    "page_number": page_index + 1,
                    "text": text,
                }
            )

    if not any(page["text"] for page in pages):
        raise ValueError(
            "No extractable text found in the PDF. "
            "It may be scanned images only (needs OCR) or empty."
        )

    return pages


def load_docx_from_bytes(docx_bytes: bytes) -> list[dict]:
    """
    Extract text from a Word document (.docx).

    Word files do not have fixed pages like PDFs, so we group paragraphs
    into pseudo-pages of roughly _DOCX_PAGE_TARGET_CHARS characters each.
    Page numbers are section indices (1, 2, 3...) for citation purposes.
    """
    if not docx_bytes:
        raise ValueError("DOCX bytes are empty.")

    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for .docx files. Install with: pip install python-docx"
        ) from exc

    doc = Document(io.BytesIO(docx_bytes))  #docx_bytes are raw bytes. ,,,io.BytesIO(...) makes them behave like a file.

    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()   #Get paragraph text and remove surrounding whitespace.
        if text:
            paragraphs.append(text)

    for table in doc.tables:     #Take all the non-empty cells in this row, clean their text, and join them together with | between them.
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    if not paragraphs:
        raise ValueError(
            "No extractable text found in the Word document. "
            "The file may be empty or contain only images."
        )

    pages: list[dict] = []   #completed pseudo-pages
    page_num = 1      #current pseudo-page number
    buffer: list[str] = []    #paragraphs currently being collected
    char_count = 0            #how many characters we've collected so far.

    for para in paragraphs:    #Take paragraphs one at a time.
        buffer.append(para)     #Put the paragraph into the current pseudo-page.
        char_count += len(para)     #Count how many characters we've collected.
        if char_count >= _DOCX_PAGE_TARGET_CHARS:   #So when we reach approximately 1500 characters: Close this pseudo-page.
            pages.append(
                {
                    "page_number": page_num,
                    "text": "\n\n".join(buffer),
                }
            )
            page_num += 1  #Move to the next pseudo-page.
            buffer = []
            char_count = 0  #Reset the character count for the next pseudo-page.

    if buffer:    #If there are any remaining paragraphs that didn't fit in the LAST pseudo-page:
        pages.append(
            {
                "page_number": page_num,
                "text": "\n\n".join(buffer),
            }
        )

    return pages


def load_document_from_bytes(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Route uploaded bytes to the correct loader based on file extension.

    Input:  raw file bytes + original filename (e.g. "report.pdf")
    Output: [{"page_number": 1, "text": "..."}, ...]
    """
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return load_pdf_from_bytes(file_bytes)
    if ext == ".docx":
        return load_docx_from_bytes(file_bytes)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
    raise ValueError(
        f"Unsupported file type '{ext or '(none)'}'. Supported types: {supported}"
    )


def clean_text(text: str) -> str:
    """Collapse runs of whitespace before chunking."""
    return " ".join(text.split())
